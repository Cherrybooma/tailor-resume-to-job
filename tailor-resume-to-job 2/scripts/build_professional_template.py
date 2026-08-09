#!/usr/bin/env python3
"""Build the polished, photo-capable Chinese resume template."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "assets" / "professional-resume-cn.docx"
NAVY = "173B57"
BLUE = "2B6F91"
TEXT = RGBColor(36, 52, 68)
MUTED = RGBColor(92, 108, 120)


def set_fonts(style, size, bold=False, color=TEXT):
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, bottom_color=None, bottom_size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "insideH", "insideV", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    bottom = OxmlElement("w:bottom")
    if bottom_color:
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), bottom_size)
        bottom.set(qn("w:color"), bottom_color)
    else:
        bottom.set(qn("w:val"), "nil")
    borders.append(bottom)


def add_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), BLUE)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D6E0E7")
    borders.extend([left, bottom])
    p_pr.append(borders)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.10)
section.bottom_margin = Cm(1.05)
section.left_margin = Cm(1.45)
section.right_margin = Cm(1.45)

normal = doc.styles["Normal"]
set_fonts(normal, 9.5)
normal.paragraph_format.space_after = Pt(1.8)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

styles = doc.styles
name = styles.add_style("Resume Name", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(name, 23, True, RGBColor(23, 59, 87))
name.paragraph_format.space_after = Pt(2)

target = styles.add_style("Resume Target", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(target, 10.5, True, RGBColor(43, 111, 145))
target.paragraph_format.space_after = Pt(3)

meta = styles.add_style("Resume Meta", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(meta, 9.1, False, MUTED)
meta.paragraph_format.space_after = Pt(0)

section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(section_style, 11.2, True, RGBColor(23, 59, 87))
section_style.paragraph_format.space_before = Pt(6)
section_style.paragraph_format.space_after = Pt(3)
section_style.paragraph_format.keep_with_next = True

entry = styles.add_style("Resume Entry", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(entry, 9.9, True, TEXT)
entry.paragraph_format.space_before = Pt(2.5)
entry.paragraph_format.space_after = Pt(1)
entry.paragraph_format.keep_with_next = True
entry.paragraph_format.tab_stops.add_tab_stop(Cm(17.9), WD_TAB_ALIGNMENT.RIGHT)

body_style = styles.add_style("Resume Body", WD_STYLE_TYPE.PARAGRAPH)
set_fonts(body_style, 9.5, False, TEXT)
body_style.paragraph_format.space_after = Pt(1.5)

bullet = styles["List Bullet"]
set_fonts(bullet, 9.5, False, TEXT)
bullet.paragraph_format.left_indent = Cm(0.45)
bullet.paragraph_format.first_line_indent = Cm(-0.25)
bullet.paragraph_format.space_after = Pt(1.2)

header = doc.add_table(rows=1, cols=2)
header.autofit = False
header.columns[0].width = Cm(15.35)
header.columns[1].width = Cm(2.45)
set_table_borders(header, "B8C8D2", "6")
left, right = header.rows[0].cells
left.width = Cm(15.35)
right.width = Cm(2.45)
left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_cell_margins(left, bottom=70, end=120)
set_cell_margins(right, bottom=70, start=80)

p = left.paragraphs[0]
p.style = "Resume Name"
p.add_run("姓名")
p = left.add_paragraph(style="Resume Target")
p.add_run("专业方向 | 核心技术能力")
p = left.add_paragraph(style="Resume Meta")
p.add_run("城市 | 手机 | 邮箱 | 个人主页/GitHub")

p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[[PHOTO]]")
r.font.size = Pt(9)
r.font.color.rgb = MUTED

for title in ("教育经历", "工作/实习经历", "项目经历", "专业技能"):
    p = doc.add_paragraph(style="Resume Section")
    p.add_run(title)
    add_paragraph_border(p)
    if title == "教育经历":
        p = doc.add_paragraph(style="Resume Entry")
        p.add_run("学校 | 学位 | 专业\t起止时间")
    elif title in ("工作/实习经历", "项目经历"):
        p = doc.add_paragraph(style="Resume Entry")
        p.add_run("单位/项目 | 职位/角色\t起止时间")
        doc.add_paragraph("使用可验证事实描述与岗位相关的行动、方法和结果。", style="List Bullet")
    else:
        doc.add_paragraph("技能类别：只保留有证据支持且与岗位相关的能力。", style="Resume Body")

doc.core_properties.title = "专业中文简历模板（支持照片）"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
