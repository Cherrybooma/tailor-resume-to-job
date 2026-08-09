#!/usr/bin/env python3
"""生成紧凑、ATS 友好的中英文简历模板。"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(31, 41, 55)
ACCENT = RGBColor(30, 64, 175)
MUTED = RGBColor(75, 85, 99)


def set_font(run, latin="Arial", east_asia="Microsoft YaHei"):
    run.font.name = latin
    props = run._element.get_or_add_rPr()
    fonts = props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_border(paragraph, color="1E40AF", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in [
        ("Resume Name", 19, INK),
        ("Resume Section", 11, ACCENT),
        ("Resume Entry", 10.2, INK),
        ("Resume Meta", 9, MUTED),
    ]:
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    styles["Resume Name"].font.bold = True
    styles["Resume Name"].paragraph_format.space_after = Pt(2)
    styles["Resume Section"].font.bold = True
    styles["Resume Section"].paragraph_format.space_before = Pt(7)
    styles["Resume Section"].paragraph_format.space_after = Pt(3)
    styles["Resume Entry"].font.bold = True
    styles["Resume Entry"].paragraph_format.space_before = Pt(2)
    styles["Resume Entry"].paragraph_format.space_after = Pt(1)
    styles["Resume Meta"].paragraph_format.space_after = Pt(2)

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.6)
    bullet.font.color.rgb = INK
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(1.2)
    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_text(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    set_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_section(doc, title):
    p = doc.add_paragraph(style="Resume Section")
    add_text(p, title, bold=True, color=ACCENT, size=11)
    set_cell_border(p)


def add_entry(doc, left, right):
    p = doc.add_paragraph(style="Resume Entry")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.7), WD_TAB_ALIGNMENT.RIGHT)
    add_text(p, left, bold=True, size=10.2)
    add_text(p, "\t" + right, bold=True, color=MUTED, size=9.3)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_text(p, text, size=9.6)


def build(path: Path, language: str):
    cn = language == "cn"
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    configure_styles(doc)

    name = "姓名" if cn else "FULL NAME"
    contact = "城市 | 手机 | 邮箱 | 个人主页/GitHub" if cn else "City | Phone | Email | Portfolio/GitHub"
    p = doc.add_paragraph(style="Resume Name")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, name, bold=True, size=19)
    p = doc.add_paragraph(style="Resume Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, contact, color=MUTED, size=9)

    if cn:
        sections = [
            ("求职方向", [("目标岗位 | 一句话证据型定位", None)]),
            ("教育经历", [("学校 | 学位 | 专业", "起止时间")]),
            ("工作/实习经历", [("公司 | 职位", "起止时间")]),
            ("项目经历", [("项目名称 | 角色", "起止时间")]),
            ("专业技能", [("编程语言：按与岗位的相关性排序；只保留有事实依据的技能。", None)]),
        ]
        bullets = [
            "使用动作 + 对象/场景 + 方法 + 已验证结果描述贡献；没有数据时不要编造数据。",
            "突出与目标 JD 最相关的职责、技术和结果，并确保每句话都可在面试中解释。",
        ]
    else:
        sections = [
            ("TARGET", [("Target role | one-line evidence-based positioning", None)]),
            ("EDUCATION", [("University | Degree | Major", "Dates")]),
            ("EXPERIENCE", [("Company | Role", "Dates")]),
            ("PROJECTS", [("Project | Role", "Dates")]),
            ("SKILLS", [("Languages: prioritize role-relevant, evidence-backed skills.", None)]),
        ]
        bullets = [
            "Describe contributions as action + context + method + verified result; never invent metrics.",
            "Prioritize responsibilities, technologies, and outcomes relevant to the target JD and defensible in an interview.",
        ]

    for title, entries in sections:
        add_section(doc, title)
        for left, right in entries:
            if right:
                add_entry(doc, left, right)
                if title in {"工作/实习经历", "项目经历", "EXPERIENCE", "PROJECTS"}:
                    for bullet in bullets:
                        add_bullet(doc, bullet)
            else:
                p = doc.add_paragraph()
                add_text(p, left, size=9.8)

    core = doc.core_properties
    core.title = "ATS Resume Template"
    core.subject = "Editable one-column resume template"
    core.author = ""
    core.last_modified_by = ""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    build(args.output_dir / "ats-resume-cn.docx", "cn")
    build(args.output_dir / "ats-resume-en.docx", "en")


if __name__ == "__main__":
    main()
